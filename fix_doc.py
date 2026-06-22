import docx
from docx.shared import Pt, Inches, Cm, RGBColor, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

INPUT = r'C:\Users\brian\Downloads\Grupo_4_TP2.docx'
OUTPUT = r'C:\Users\brian\Downloads\Grupo_4_TP2_ARREGLADO.docx'

doc = docx.Document(INPUT)

# ======================================================
# 1. Collect headings
# ======================================================
headings = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    
    is_heading = False
    for r in p.runs:
        if r.bold and r.font.size and r.font.size >= Pt(14):
            is_heading = True
            break
    
    if text in ['ÍNDICE', 'Pol\u00edtica de Calidad del Proyecto',
                'RELEVAMIENTO DE DOCUMENTACI\u00d3N', 'Roles del equipo',
                'Roadmap']:
        is_heading = True
    if re.match(r'^\d+\.\s', text) and any(r.bold for r in p.runs if r.text.strip()):
        is_heading = True
    if re.match(r'^\d+\.\d+\s', text):
        is_heading = True
    
    if is_heading:
        level = 1 if re.match(r'^\d+\.\d+\s', text) else 0
        # Clean up: take only first line / sentence for the index
        clean_text = text.split('\n')[0].strip()
        # Remove parenthetical notes from index
        clean_text = re.sub(r'\s*\(.*?\)\s*$', '', clean_text).strip()
        # Skip informal/draft headings
        skip_keywords = ['esto hay que mostrar', 'DURMIENDO', 'CS', 'LOL', 'joda',
                         'Informe de Avance Semanal']
        if any(kw.lower() in clean_text.lower() for kw in skip_keywords):
            continue
        headings.append((clean_text[:80], level, i))

# ======================================================
# 2. Write detailed index in paragraphs [5] through [22]
# ======================================================
index_items = []
for h, lvl, idx in headings:
    indent = '    ' if lvl > 0 else ''
    bullet = '-' if lvl > 0 else '*'
    index_items.append(f'{indent}{bullet} {h}')

# Paragraph [5] -> "ÍNDICE" title
p = doc.paragraphs[5]
for r in p.runs:
    r.text = ''
p.runs[0].text = 'INDICE'
p.runs[0].bold = True
p.runs[0].font.size = Pt(18)
p.runs[0].font.name = 'Calibri'
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Fill paragraphs [6]..[22] with index items
for offset, item in enumerate(index_items):
    idx = 6 + offset
    if idx >= 23:
        break
    p = doc.paragraphs[idx]
    for r in p.runs:
        r.text = ''
    if not p.runs:
        p.add_run('')
    p.runs[0].text = item
    p.runs[0].bold = False
    p.runs[0].font.size = Pt(12) if item.startswith(' ') else Pt(13)
    p.runs[0].font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

print(f'Index written: {len(index_items)} items')

# ======================================================
# 3. Fix tables
# ======================================================
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    
    # Remove old borders if any
    old_borders = tblPr.findall(qn('w:tblBorders'))
    for ob in old_borders:
        tblPr.remove(ob)
    
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Remove old vAlign
            old_valign = tcPr.findall(qn('w:vAlign'))
            for ov in old_valign:
                tcPr.remove(ov)
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            tcPr.append(vAlign)
            
            # Remove old shading then add new
            old_shd = tcPr.findall(qn('w:shd'))
            for os in old_shd:
                tcPr.remove(os)
            if ri == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                tcPr.append(shading)
            
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    if ri == 0:
                        r.bold = True
                        r.font.size = Pt(10)
                    else:
                        if r.font.size is None or r.font.size > Pt(10):
                            r.font.size = Pt(9)
                    r.font.name = 'Calibri'

print(f'Tables formatted: {len(doc.tables)} tables')

# ======================================================
# 4. Fix paragraph alignment
# ======================================================
heading_texts = {'\u00cdNDICE', 'Pol\u00edtica de Calidad del Proyecto',
                 'RELEVAMIENTO DE DOCUMENTACI\u00d3N', 'Roles del equipo',
                 'Roadmap', 'Informe de Avance Semanal'}

for p in doc.paragraphs:
    text = p.text.strip()
    if not text:
        continue
    
    is_heading = False
    for r in p.runs:
        if r.bold and r.font.size and r.font.size >= Pt(14):
            is_heading = True
            break
    
    if text in heading_texts:
        is_heading = True
    if re.match(r'^\d+\.\s', text) or re.match(r'^\d+\.\d+\s', text):
        is_heading = True
    
    if is_heading:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

print('Text alignment fixed')

# ======================================================
# 5. Save
# ======================================================
doc.save(OUTPUT)
print(f'\nSaved to: {OUTPUT}')
