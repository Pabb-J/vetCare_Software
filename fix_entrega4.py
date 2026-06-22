import docx
from docx.shared import Pt, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re

INPUT = r'C:\Users\brian\Downloads\Grupo_4_TP2_entrega_4.docx'
OUTPUT = r'C:\Users\brian\Downloads\Grupo_4_TP2_entrega_4_ARREGLADO.docx'

doc = docx.Document(INPUT)

# ======================================================
# 1. FIX INDEX AND SPECIAL HEADERS (paragraphs 1-42)
# ======================================================
print('Fixing index and headers...')

# Fix "Indice" title (paragraph 1)
p = doc.paragraphs[1]
if 'Indice' in p.text:
    for r in p.runs:
        r.font.bold = True
        r.font.size = Pt(16)
        r.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Fix paragraphs that are RIGHT-aligned (like "Evaluacion CMMI", "Propuesta Scrum")
for p in doc.paragraphs:
    if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs:
            if r.font.name is None:
                r.font.name = 'Calibri'

# ======================================================
# 2. FIX ALL TABLES: borders, shading, alignment
# ======================================================
print(f'Fixing {len(doc.tables)} tables...')
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    
    # Remove old borders, add new ones
    for ob in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(ob)
    
    borders = parse_xml(
        '<w:tblBorders %s>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Vertical center
            for ov in tcPr.findall(qn('w:vAlign')):
                tcPr.remove(ov)
            vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
            tcPr.append(vAlign)
            
            # Header row shading
            for os in tcPr.findall(qn('w:shd')):
                tcPr.remove(os)
            if ri == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
                tcPr.append(shading)
            
            for cp in cell.paragraphs:
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in cp.runs:
                    if ri == 0:
                        r.font.bold = True
                        r.font.size = Pt(10)
                    r.font.name = 'Calibri'
                    if r.font.size is None:
                        r.font.size = Pt(9)

# ======================================================
# 3. FIX TEXT ALIGNMENT
# ======================================================
print('Fixing text alignment...')

# These are heading texts that should stay LEFT-aligned
heading_list = [
    'Indice', 'INDICE',
    'Sistema Integral V', 'INTRODUCCION', 'SELECCION Y DESCRIPCION',
    'Nombre y descripcion', 'Materia y a', 'Tecnolog',
    'Funcionalidades principales', 'Problemas y limitaciones',
    'Relevamiento de Documentaci', 'Modelado y Dise',
    'Estrategia General de Calidad', 'Objetivo del An',
    'Normas Seleccionadas', 'Divisi', 'Cronograma de trabajo',
    'L', 'Organizaci', 'Puntos de mejora', 'Propuesta de mejora',
    'Pol', 'ANALISTA DE CALIDAD DE PRODUCTO', 'Evaluaci',
    'Analista de Procesos',
    'Especialista en Testing', 'Mini Plan de SQA',
    'Casos de prueba', 'Analista de Metodolog',
    'Evaluaci', 'Propuesta Scrum', 'Conclusi',
    'INFORME DE AVANCE', 'Alcance del An',
    'Caracter', 'Acciones de mejora',
    'Reconstrucci', 'Comparaci', 'Riesgos detectados',
    'Propuesta', '1. Introducci', '2. Resumen Ejecutivo',
    '3. Evaluaci', 'Seguridad', 'Confiabilidad',
    'Usabilidad', 'Mantenibilidad',
    # Extra index items
    'Introduccion', 'Seleccion y Descripcion',
    'Nombre y descripcion general', 'Materia y ano',
    'Tecnologias utilizadas',
]

def strip_accents(s):
    import unicodedata
    return ''.join(c for c in unicodedata.normalize('NFKD', s) if not unicodedata.combining(c))

def is_heading_paragraph(p):
    t = p.text.strip()
    if not t:
        return False
    
    # Check by text content (normalize accents for comparison)
    t_norm = strip_accents(t.upper())
    for h in heading_list:
        h_norm = strip_accents(h.upper())
        if t_norm.startswith(h_norm):
            return True
    
    # Check by formatting (bold + large)
    for r in p.runs:
        if r.text.strip() and r.font.bold:
            if r.font.size and r.font.size >= Pt(13):
                return True
            # Also check if it's short and bold
            if len(t) < 60:
                return True
    
    # Numbered sections
    if re.match(r'^\d+\.\s', t) or re.match(r'^\d+\.\d+\s', t):
        return True
    
    return False

for p in doc.paragraphs:
    t = p.text.strip()
    if not t:
        continue
    
    if is_heading_paragraph(p):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ======================================================
# 4. SAVE
# ======================================================
doc.save(OUTPUT)
print(f'Saved to: {OUTPUT}')
print('Done!')
