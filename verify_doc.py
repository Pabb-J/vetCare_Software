import docx

doc = docx.Document(r'C:\Users\brian\Downloads\Grupo_4_TP2_ARREGLADO.docx')

print('=== INDEX VERIFICATION ===')
for i in range(5, 23):
    p = doc.paragraphs[i]
    t = p.text.strip()
    if t:
        print(f'  [{i}] {t[:120]}')

print()
print('=== TABLE VERIFICATION ===')
for ti, table in enumerate(doc.tables):
    print(f'Table {ti}: {len(table.rows)} rows')
    for ri, row in enumerate(table.rows):
        c0 = row.cells[0].text.strip()[:30]
        print(f'  Row {ri}: first cell = "{c0}"')

print()
print('=== ALIGNMENT SAMPLES ===')
for i in [23, 24, 29, 30, 58, 59, 74, 75, 96, 100]:
    p = doc.paragraphs[i]
    t = p.text.strip()[:80]
    a = str(p.alignment) if p.alignment is not None else 'None'
    print(f'  [{i:3d}] align={a:15s} | {t}')
