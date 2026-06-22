import docx
from docx.shared import Pt

doc = docx.Document(r'C:\Users\brian\Downloads\Grupo_4_TP2_entrega_4_ARREGLADO.docx')

print('=== INDEX PARAGRAPHS (1-40) ===')
for i in range(0, 41):
    if i >= len(doc.paragraphs):
        break
    p = doc.paragraphs[i]
    t = p.text.strip()
    a = str(p.alignment) if p.alignment is not None else 'None'
    if t:
        print(f'  [{i:3d}] align={a:15s} | {t[:120]}')

print()
print('=== ALL TABLES ===')
for ti, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    # Check borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    has_borders = len(tblPr.findall(qn('w:tblBorders'))) > 0 if tblPr is not None else False
    first_cell = table.rows[0].cells[0].text.strip()[:40] if rows > 0 else ''
    print(f'  Table {ti:2d}: {rows}r x {cols}c  borders={has_borders}  first_cell="{first_cell}"')
    # Check alignment of last row
    if rows > 1:
        last_row_align = table.rows[-1].cells[0].paragraphs[0].alignment if table.rows[-1].cells[0].paragraphs else None
        print(f'           last row align={last_row_align}')
